"""Build a KBID proposal .docx that reproduces KBID's REAL template.

Layout reverse-engineered from the sample PDFs (Info Sheet/): US Letter, 1" margins,
Arial throughout; KBID K logomark on the cover; running "KBID" footer + page number;
large light page titles; bold subheadings (no rules); justified body; fee blocks as
an indented two-column tabbed list with an underlined mini-heading, a thin rule, and a
bold total; Word bullet lists; a two-column signature block. NOT the redesign.
"""
from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ASSETS = Path(__file__).resolve().parent.parent / "assets"
INK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x6F, 0x68, 0x58)
FONT = "Arial"
RIGHT_TAB = Inches(6.5)          # amount column tab (within 6.5" text width)
FEE_INDENT = Inches(1.0)


def _money(n):
    try:
        return "$" + format(int(round(float(n))), ",")
    except Exception:
        return str(n)


def _run(p, text, *, bold=False, italic=False, size=11, color=INK, underline=False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.underline = underline
    r.font.color.rgb = color
    return r


def _para(doc, space_after=6, space_before=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    return p


def _body(doc, text, justify=True):
    p = _para(doc, space_after=10)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, text, size=11)
    return p


def _subhead(doc, text):
    p = _para(doc, space_after=3, space_before=10)
    _run(p, text, bold=True, size=12)
    return p


def _title(doc, text):
    p = _para(doc, space_after=10, space_before=0)
    _run(p, text, size=22)
    return p


def _bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        _run(p, it, size=11)


def _fee_line(doc, label, amount, *, bold=False, underline_label=False, rule=False, indent=FEE_INDENT):
    p = _para(doc, space_after=1)
    p.paragraph_format.left_indent = indent
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)
    _run(p, label, bold=bold, underline=underline_label, size=11)
    if amount is not None:
        _run(p, "\t" + amount, bold=bold, size=11)
    if rule:
        _bottom_border(p)
    return p


def _mini_heading(doc, text):
    p = _para(doc, space_after=2, space_before=8)
    p.paragraph_format.left_indent = FEE_INDENT
    _run(p, text, bold=True, underline=True, size=11)
    return p


def _bottom_border(p):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom); pPr.append(pbdr)


def _page_number(run):
    fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1); run._r.append(instr); run._r.append(fldChar2)


def _footer(section, firm):
    f = section.footer
    f.is_linked_to_previous = False
    p = f.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    _run(p, "KBID ", bold=True, size=9)
    _run(p, firm["legal"], size=9, color=GRAY)
    _run(p, "\t", size=9)
    _page_number(_run(p, "", size=9, color=GRAY))


def _header(section, text):
    h = section.header
    h.is_linked_to_previous = False
    p = h.paragraphs[0]
    _run(p, text, size=9, color=GRAY)


def _fees_block(doc, sec):
    est = sec.get("estimate") or {}
    phases = est.get("phases") or []
    mode = sec.get("fee_mode", "lump")
    cons = {k: v for k, v in (sec.get("consultants") or {}).items() if v not in ("", None)}
    _subhead(doc, "Fees & Reimbursable Expenses")
    if mode == "hourly":
        _body(doc, "The fees based on the interior design services to be performed and number of meetings per this "
                   "Letter Agreement will be billed at an hourly rate and only for hours completed according to the "
                   "scope options listed above. Invoices shall be due monthly at the following hourly rates.")
        _mini_heading(doc, "Estimated Hourly Fee Range")
        lo_sum = hi_sum = 0
        for p in phases:
            base = p.get("total_rounded") or p.get("total_raw") or 0
            lo = int(round(base / 1000.0) * 1000); hi = lo + 2000
            lo_sum += lo; hi_sum += hi
            hrs = int(round(p.get("hours") or 0)); wk = _fmt_num(p.get("duration_weeks"))
            _fee_line(doc, f"{p['name']} ({hrs} Hours / {wk} weeks)", f"{_money(lo)} - {_money(hi)}")
        _fee_line(doc, "Construction Observation", "Hourly As Needed", rule=True)
        _fee_line(doc, "Total", f"{_money(lo_sum)} - {_money(hi_sum)}", bold=True)
        if cons.get("Structural"):
            _para(doc, space_after=1)
            _fee_line(doc, "Allowance for Structural Engineer", _money(cons["Structural"]))
    else:
        _body(doc, "The fees based on the interior design services to be performed and the set hours and number of "
                   "meetings per this Letter Agreement shall be due on a lump-sum basis per design process stage as follows:")
        _mini_heading(doc, "Interior Design Fee Breakdown")
        total = 0
        last_alloc = None
        for p in phases:
            amt = p.get("total_rounded") or p.get("total_raw") or 0
            total += amt
            _fee_line(doc, p["name"], _money(amt))
        for label, key in (("MEP Engineering Fee Allowance", "MEP"), ("Structural Engineering Allowance", "Structural")):
            if cons.get(key):
                total += float(cons[key]); last_alloc = label
                _fee_line(doc, label, _money(cons[key]), rule=(label == last_alloc))
        _fee_line(doc, "Total Design Fee", _money(total), bold=True, rule=(last_alloc is None))
    # rate table
    _mini_heading(doc, "Interior Design Rate Table")
    for label, rate in (("Owner", 210), ("Design Director", 200), ("Interior Designer", 170), ("Designer", 150)):
        _fee_line(doc, label, f"${rate} / hr")


def _fmt_num(n):
    try:
        f = float(n)
        return str(int(f)) if f == int(f) else str(f)
    except Exception:
        return str(n)


def build_docx(sec: dict) -> bytes:
    meta = sec.get("meta", {})
    firm = meta.get("firm", {})
    doc = Document()
    # base style
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    for m in ("top", "bottom", "left", "right"):
        setattr(section, f"{m}_margin", Inches(1))
    _footer(section, firm)

    # ---- cover ----
    logo = ASSETS / "kbid_logo.png"
    if logo.exists():
        doc.add_picture(str(logo), width=Inches(0.55))
    _para(doc, space_after=18)
    # designer block
    p = _para(doc, space_after=0); _run(p, f"DESIGNER: {firm.get('name','')}", bold=True)
    for line in (firm.get("address"), firm.get("city"), firm.get("email"), firm.get("phone")):
        pp = _para(doc, space_after=0); _run(pp, line or "", color=GRAY)
    _para(doc, space_after=10)
    # client block
    p = _para(doc, space_after=0); _run(p, f"CLIENT: {meta.get('client_name','')}", bold=True)
    for line in (meta.get("client_address"), meta.get("client_phone"), meta.get("client_email")):
        if line:
            pp = _para(doc, space_after=0); _run(pp, line, color=GRAY)
    _para(doc, space_after=12)
    p = _para(doc, space_after=10); _run(p, meta.get("date", ""), bold=True)
    p = _para(doc, space_after=0); _run(p, f"RE: {meta.get('re','')}", bold=True)
    if meta.get("project_address"):
        pp = _para(doc, space_after=10); pp.paragraph_format.left_indent = Inches(0.3); _run(pp, meta["project_address"])
    p = _para(doc, space_after=10); _run(p, f"Dear {meta.get('client_contact') or meta.get('client_name','')},")
    _body(doc, sec.get("cover_opener", ""))
    for para in (sec.get("cover_body") or []):
        _body(doc, para)
    _para(doc, space_after=24)
    p = _para(doc, space_after=2); _run(p, "Sincerely,")
    _para(doc, space_after=18)
    p = _para(doc, space_after=0)
    _run(p, "Kali Buchanan, ", size=11)
    _run(p, "RID, NCIDQ, IIDA, LEED AP", size=9)

    # ---- interior ----
    doc.add_section(WD_SECTION.NEW_PAGE)
    interior = doc.sections[-1]
    interior.page_width = Inches(8.5); interior.page_height = Inches(11)
    for m in ("top", "bottom", "left", "right"):
        setattr(interior, f"{m}_margin", Inches(1))
    _footer(interior, firm)
    _header(interior, meta.get("project_name", ""))

    _title(doc, "Project Scope & Fees")
    if sec.get("scope_overview"):
        _body(doc, sec["scope_overview"])
    for ph in sec.get("phases", []):
        _subhead(doc, ph["name"]); _body(doc, ph["description"])
    _subhead(doc, "Additional Terms & Conditions"); _body(doc, sec.get("terms", ""))
    _subhead(doc, "Exclusions, Qualifications, or Exceptions")
    _body(doc, "KBID excludes from its services as well as states as qualifications and exceptions the following.", justify=False)
    _bullets(doc, sec.get("exclusions", []))
    if sec.get("schedule"):
        _subhead(doc, "Schedule"); _body(doc, sec["schedule"])
    _fees_block(doc, sec)
    _subhead(doc, "Payment"); _body(doc, sec.get("payment", ""))
    _subhead(doc, "Reimbursable Expenses")
    _body(doc, "The following are typical reimbursable expenses incurred by KBID that are not included in the fees "
               "above and will be included in the monthly invoicing in addition to the Interior Design Fees.", justify=False)
    _bullets(doc, sec.get("reimbursables", []))
    _subhead(doc, "Documents & Photographs"); _body(doc, sec.get("documents", ""))
    _subhead(doc, "Limitation on Liability"); _body(doc, sec.get("limitation", ""))
    _subhead(doc, "Notice to Proceed"); _body(doc, sec.get("notice", ""))
    if sec.get("cover_closer"):
        _body(doc, sec["cover_closer"])

    # ---- signature block (two columns) ----
    _para(doc, space_after=18)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = True
    left, right = tbl.rows[0].cells
    def sig(cell, party, name, title):
        cell.paragraphs[0].text = ""
        _run(cell.paragraphs[0], party, bold=True)
        _run(cell.add_paragraph(), "By: ______________________________")
        _run(cell.add_paragraph(), name)
        _run(cell.add_paragraph(), title, size=9, color=GRAY)
        _run(cell.add_paragraph(), "Date: ____________________", size=10)
    sig(left, 'DESIGNER: Kali Buchanan Interior Design "KBID"', firm.get("signatory", ""), "Owner, Officer for KBID")
    sig(right, f"CLIENT: {meta.get('client_name','')}", meta.get("client_contact") or meta.get("client_name", ""), meta.get("client_title", "") or "Client")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
